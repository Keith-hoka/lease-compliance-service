import io
from collections import Counter
from pathlib import Path

import pytest
from docx import Document

from app.ingest import parser_vic
from app.ingest.parser_vic import SECTION_HEADING_STYLES, parse_docx

HEAD = SECTION_HEADING_STYLES[0]


def build_docx(paragraphs: list[tuple[str | None, str]]) -> bytes:
    """A minimal DOCX from (style_name, text) rows; None = default style."""
    doc = Document()
    for style_name, text in paragraphs:
        paragraph = doc.add_paragraph(text)
        if style_name is not None:
            from docx.enum.style import WD_STYLE_TYPE

            styles = doc.styles
            if style_name not in [s.name for s in styles]:
                styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            paragraph.style = styles[style_name]
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def test_sections_split_with_part_and_division_tracking():
    data = build_docx(
        [
            (None, "Part 2—Tenancy agreements"),
            (None, "Division 1—General"),
            (HEAD, "26\tApplication of Part"),
            (None, "This Part applies to all agreements."),
            (HEAD, "27B\tProhibited terms"),
            (None, "A term listed below must not be included."),
            (None, "Penalty: 60 penalty units."),
        ]
    )
    sections = parse_docx(data)
    assert [s.section_no for s in sections] == ["26", "27B"]
    assert sections[1].heading == "Prohibited terms"
    assert "must not be included" in sections[1].body_text
    assert "Penalty" in sections[1].body_text
    assert sections[0].part == "Part 2—Tenancy agreements"
    assert sections[0].division == "Division 1—General"


def test_subdivision_updates_division_label():
    data = build_docx(
        [
            (None, "Part 2—Tenancy agreements"),
            (None, "Division 2—Applications"),
            (HEAD, "10\tFirst"),
            (None, "Body one."),
            (None, "Subdivision 1—Application to rental agreements"),
            (HEAD, "11\tSecond"),
            (None, "Body two."),
        ]
    )
    sections = parse_docx(data)
    assert sections[0].division == "Division 2—Applications"
    assert sections[1].division == "Subdivision 1—Application to rental agreements"


def test_body_lines_starting_with_numbers_stay_in_body():
    data = build_docx(
        [
            (None, "Part 1—Preliminary"),
            (HEAD, "5\tDefinitions"),
            (None, "14 days means a fortnight in this test body."),
        ]
    )
    sections = parse_docx(data)
    assert [s.section_no for s in sections] == ["5"]
    assert "14 days" in sections[0].body_text


def test_regex_fallback_when_styles_unpinned(monkeypatch):
    from app.ingest import parser_vic

    monkeypatch.setattr(parser_vic, "SECTION_HEADING_STYLES", ())
    data = build_docx(
        [
            (None, "Part 1—Preliminary"),
            (None, "1 Purposes"),
            (None, "The purposes of this Act are set out."),
        ]
    )
    sections = parse_docx(data)
    assert [s.section_no for s in sections] == ["1"]


def test_toc_lines_are_skipped_by_style():
    data = build_docx(
        [
            ("toc 3", "27B Prohibited terms 55"),
            (None, "Part 1—Preliminary"),
            (HEAD, "1\tPurposes"),
            (None, "The purposes of this Act are set out."),
        ]
    )
    sections = parse_docx(data)
    assert [s.section_no for s in sections] == ["1"]


def test_parsing_stops_at_endnotes():
    data = build_docx(
        [
            (None, "Part 1—Preliminary"),
            (HEAD, "1\tPurposes"),
            (None, "Body text."),
            (None, "Endnotes"),
            (HEAD, "2\tThis looks like a section but is history."),
        ]
    )
    sections = parse_docx(data)
    assert [s.section_no for s in sections] == ["1"]


def test_schedule_clauses_get_prefixed_keys():
    data = build_docx(
        [
            (None, "Part 1—Preliminary"),
            (HEAD, "1\tPurposes"),
            (None, "Body."),
            (None, "Schedule 1—Transitional provisions"),
            (HEAD, "1\tSaved instruments"),
            (None, "Schedule body."),
            (None, "Schedule 1A—Pecuniary penalty provisions"),
            (HEAD, "2\tPenalty items"),
            (None, "Item list."),
            (None, "Schedule 4—Validation"),
            (None, "Part 1—Preliminary"),
            (HEAD, "1\tDefinition"),
            (None, "Internal part body."),
        ]
    )
    sections = parse_docx(data)
    assert [s.section_no for s in sections] == ["1", "S1-1", "S1A-2", "S4-1"]
    assert sections[1].part == "Schedule 1—Transitional provisions"
    assert sections[1].division is None
    assert sections[3].part == "Schedule 4—Validation"
    assert sections[3].division == "Part 1—Preliminary"


def test_repealed_placeholder_loads_as_shown():
    data = build_docx(
        [
            (None, "Part 1—Preliminary"),
            (HEAD, "27A\tRepealed"),
            (None, "* * * * *"),
        ]
    )
    sections = parse_docx(data)
    assert sections[0].section_no == "27A"
    assert sections[0].heading == "Repealed"


def test_partless_document_collects_sections_with_pinned_styles():
    """The Regulations open with clauses before any Part heading."""
    data = build_docx(
        [
            (None, "Version No. 009"),
            (HEAD, "1\tObjective"),
            (None, "The objective of these Regulations is to prescribe matters."),
            (HEAD, "2\tAuthorising provision"),
            (None, "Made under section 511."),
        ]
    )
    sections = parse_docx(data)
    assert [s.section_no for s in sections] == ["1", "2"]
    assert sections[0].part is None


def test_form_terms_parse_with_form_scoped_keys():
    data = build_docx(
        [
            ("Heading - PART", "Schedule 1—Forms"),
            ("Side Note", "Sch. 1 Form 1 amended by S.R. No. 123/2025."),
            ("New Form Heading", "Form 1"),
            (None, "Residential Tenancies Act 1997"),
            ("New Form Heading", "Residential rental agreement of no more than 5 years"),
            ("New Form Heading", "PART A—GENERAL"),
            (None, "1.\tDate of agreement"),
            (None, "This is the date the agreement is signed."),
            ("Side Note", "amendment note inside the form"),
            (None, "2.\tPremises"),
            (None, "Address of premises."),
            ("New Form Heading", "PART B—Standard Terms"),
            (None, "3.\tRent"),
            (None, "Rent must be paid on time."),
            ("New Form Heading", "Form 2"),
            ("New Form Heading", "Agreement of more than 5 years"),
            (None, "1.\tDate of agreement"),
            (None, "Second form first term."),
        ]
    )
    sections = parse_docx(data)
    by_no = {s.section_no: s for s in sections}
    assert set(by_no) == {"S1-F1-T1", "S1-F1-T2", "S1-F1-T3", "S1-F2-T1"}

    first = by_no["S1-F1-T1"]
    assert first.heading == "Date of agreement"
    assert first.body_text == "This is the date the agreement is signed."
    assert first.part == "Schedule 1—Forms"
    assert first.division == "Form 1 Residential rental agreement of no more than 5 years"

    rent = by_no["S1-F1-T3"]
    assert rent.heading == "Rent"
    assert rent.body_text == "Rent must be paid on time."

    second_form = by_no["S1-F2-T1"]
    assert second_form.division == "Form 2 Agreement of more than 5 years"
    assert second_form.body_text == "Second form first term."


def test_form_term_matches_dot_space_tab_convention():
    """The source uses two numbering conventions for top-level form
    items: dot-tab (Forms 1-7, 16A) and dot-space-tab (184 items across
    the cache, most of Forms 8-24, e.g. Form 11). A digit sub-item like
    "9.1\tx" must stay body text, not start a new term."""
    data = build_docx(
        [
            ("Heading - PART", "Schedule 1—Forms"),
            ("New Form Heading", "Form 11"),
            ("New Form Heading", "A form title"),
            (None, "1. \tFirst term"),
            (None, "First body."),
            (None, "9.1\tSub-item text that must stay in the body."),
            (None, "2. \tSecond term"),
            (None, "Second body."),
        ]
    )
    sections = parse_docx(data)
    by_no = {s.section_no: s for s in sections}
    assert set(by_no) == {"S1-F11-T1", "S1-F11-T2"}
    assert by_no["S1-F11-T1"].heading == "First term"
    assert "Sub-item text that must stay in the body" in by_no["S1-F11-T1"].body_text
    assert by_no["S1-F11-T2"].heading == "Second term"
    assert by_no["S1-F11-T2"].body_text == "Second body."


def test_form_terms_coexist_with_schedule_clauses():
    data = build_docx(
        [
            ("Draft Heading 1", "12 Body clause"),
            (None, "Body clause text."),
            ("Heading - PART", "Schedule 1—Forms"),
            ("Draft Heading 1", "5 Schedule clause"),
            (None, "Schedule clause text."),
            ("New Form Heading", "Form 1"),
            ("New Form Heading", "A form title"),
            (None, "1.\tOnly term"),
            (None, "Term body."),
        ]
    )
    sections = parse_docx(data)
    by_no = {s.section_no: s for s in sections}
    assert set(by_no) == {"12", "S1-5", "S1-F1-T1"}
    assert by_no["12"].body_text == "Body clause text."
    assert by_no["S1-5"].body_text == "Schedule clause text."


def test_form_opener_recognised_when_styled_normal():
    """Some cached versions style a form's opener paragraph as Normal
    instead of New Form Heading; a missed opener must not fold the next
    form's terms into the previous form's keys."""
    data = build_docx(
        [
            ("Heading - PART", "Schedule 1—Forms"),
            ("New Form Heading", "Form 4"),
            ("New Form Heading", "Condition report"),
            (None, "1.\tCondition item"),
            (None, "Condition body."),
            (None, "Form 5"),
            ("New Form Heading", "Notice of proposed rent increase"),
            (None, "1.\tNotice item"),
            (None, "Notice body."),
        ]
    )
    sections = parse_docx(data)
    section_nos = [s.section_no for s in sections]
    assert set(section_nos) == {"S1-F4-T1", "S1-F5-T1"}
    assert len(section_nos) == 2

    by_no = {s.section_no: s for s in sections}
    assert by_no["S1-F4-T1"].heading == "Condition item"
    assert by_no["S1-F4-T1"].body_text == "Condition body."
    assert by_no["S1-F4-T1"].division == "Form 4 Condition report"
    assert by_no["S1-F5-T1"].heading == "Notice item"
    assert by_no["S1-F5-T1"].body_text == "Notice body."
    assert by_no["S1-F5-T1"].division == "Form 5 Notice of proposed rent increase"


def test_form_opener_recognised_when_all_caps():
    """Cached versions 005+ carry all-caps openers ("FORM 3A", "FORM 16A")
    in New Form Heading style, rescued by _FORM_RE's IGNORECASE; 001 has a
    Normal-styled "Form 5", rescued by the exact-text fallback. This pins
    the synthetic combination (Normal AND all-caps) so both matchers stay
    case-insensitive."""
    data = build_docx(
        [
            ("Heading - PART", "Schedule 1—Forms"),
            ("New Form Heading", "Form 8"),
            ("New Form Heading", "Condition report"),
            (None, "1.\tCondition item"),
            (None, "Condition body."),
            (None, "FORM 9"),
            ("New Form Heading", "Notice of proposed rent increase"),
            (None, "1.\tNotice item"),
            (None, "Notice body."),
        ]
    )
    sections = parse_docx(data)
    section_nos = [s.section_no for s in sections]
    assert set(section_nos) == {"S1-F8-T1", "S1-F9-T1"}
    assert len(section_nos) == 2

    by_no = {s.section_no: s for s in sections}
    assert by_no["S1-F9-T1"].heading == "Notice item"
    assert by_no["S1-F9-T1"].division == "Form 9 Notice of proposed rent increase"


def test_form_with_restarting_numbers_is_skipped_entirely():
    """A form whose internal numbering restarts (Form 3A's three
    independently-numbered PART sequences in real cached versions) is
    unparseable under the continuous-numbering model: it must yield
    zero terms, while a well-numbered form that follows still parses."""
    data = build_docx(
        [
            ("Heading - PART", "Schedule 1—Forms"),
            ("New Form Heading", "Form 3A"),
            ("New Form Heading", "Site agreement"),
            (None, "1.\tFirst item"),
            (None, "First body."),
            (None, "2.\tSecond item"),
            (None, "Second body."),
            (None, "1.\tRestarted item"),
            (None, "Restarted body."),
            ("New Form Heading", "Form 4"),
            ("New Form Heading", "Condition report"),
            (None, "1.\tCondition item"),
            (None, "Condition body."),
        ]
    )
    sections = parse_docx(data)
    section_nos = {s.section_no for s in sections}
    assert not any(no.startswith("S1-F3A-") for no in section_nos)
    assert section_nos == {"S1-F4-T1"}


def test_long_after_tab_text_becomes_body_not_heading():
    """A form term whose after-tab text is long, free-flowing prose
    (no short title) must not have it truncated into `heading`; it
    becomes an empty heading with the prose as the first body line,
    with nothing lost."""
    long_text = (
        "A rental provider must not request or receive a payment of rent "
        "more than 14 days in advance of the day on which the rent is "
        "payable under the rental agreement, except as otherwise agreed "
        "in writing by the renter."
    )
    assert len(long_text) > 150
    data = build_docx(
        [
            ("Heading - PART", "Schedule 1—Forms"),
            ("New Form Heading", "Form 3"),
            ("New Form Heading", "Information for renters"),
            (None, f"1.\t{long_text}"),
            (None, "Continuation paragraph."),
        ]
    )
    sections = parse_docx(data)
    assert len(sections) == 1
    term = sections[0]
    assert term.heading == ""
    assert term.body_text == f"{long_text} Continuation paragraph."


def test_numbered_body_lines_do_not_become_terms_outside_forms():
    data = build_docx(
        [
            ("Draft Heading 1", "12 Body clause"),
            (None, "1.\tThis is body prose with a tab, not a form term."),
        ]
    )
    sections = parse_docx(data)
    assert [s.section_no for s in sections] == ["12"]
    assert "body prose" in sections[0].body_text


REGS_CACHE = Path("data/raw/vic/residential-tenancies-regulations-2021")


def test_real_regs_cache_yields_form_terms():
    cached = sorted(REGS_CACHE.glob("*.docx"))
    if not cached:
        pytest.skip("VIC regulations cache not present")
    sections = parse_docx(cached[-1].read_bytes())
    form_terms = [s for s in sections if "-F" in s.section_no]
    f1 = [s for s in form_terms if s.section_no.startswith("S1-F1-")]
    f2 = [s for s in form_terms if s.section_no.startswith("S1-F2-")]
    # Exact counts drift with amendments; floors match the probed current
    # version (F1=32, F2=40, 403 total after Form 3A's deliberate skip).
    assert len(f1) >= 30
    assert len(f2) >= 38
    assert len(form_terms) >= 400
    # Guard against a form silently vanishing to a future false restart:
    # the newest version yields exactly 25 form prefixes (every form with
    # numbered items, minus Form 3A's deliberate restart skip).
    form_prefixes = {s.section_no.split("-")[1] for s in form_terms}
    assert len(form_prefixes) == 25
    assert any(s.heading == "Rent" for s in f1)
    assert all(s.part == "Schedule 1—Forms" for s in form_terms)
    schedule_clauses = [
        s for s in sections if s.section_no.startswith("S") and "-F" not in s.section_no
    ]
    assert len(schedule_clauses) >= 35


def _expected_form_term_count(data: bytes) -> int:
    """Independent scan for the Important-2 completeness check: track
    form scope with the parser's own opener predicate, count raw-text
    term matches while a form is open, and drop the forms the parser
    deliberately discards (Form 3A's restart)."""
    document = Document(io.BytesIO(data))
    in_schedule, form_no = False, None
    per_form: Counter[str] = Counter()
    for paragraph in document.paragraphs:
        style_name = paragraph.style.name if paragraph.style is not None else ""
        if style_name.lower().startswith("toc"):
            continue
        text = parser_vic._clean(paragraph.text)
        if not text:
            continue
        if text == "Endnotes":
            break
        if parser_vic._SCHEDULE_RE.match(text):
            in_schedule, form_no = True, None
            continue
        if in_schedule and (
            style_name == parser_vic.FORM_HEADING_STYLE or parser_vic._FORM_RE.fullmatch(text)
        ):
            form_match = parser_vic._FORM_RE.match(text)
            form_no = form_match.group(1).upper() if form_match else form_no
            continue
        if form_no is not None and parser_vic._FORM_TERM_RE.match(paragraph.text):
            per_form[form_no] += 1
    return sum(count for form, count in per_form.items() if form != "3A")


def test_all_cached_regs_versions_have_unique_section_numbers():
    """Pins historical-version robustness: every cached version must
    parse with no duplicate section_no and no oversized heading or
    division. Every cached version styles the Form 5 opener as Normal
    instead of New Form Heading, so this also exercises the exact-text
    opener fallback. Also cross-checks the parsed form-term count
    against an independent raw-paragraph scan, so a dropped item can't
    silently survive because the parser agrees with itself."""
    cached = sorted(REGS_CACHE.glob("*.docx"))
    if not cached:
        pytest.skip("VIC regulations cache not present")
    for path in cached:
        data = path.read_bytes()
        sections = parse_docx(data)
        counts = Counter(s.section_no for s in sections)
        _, most_common_count = counts.most_common(1)[0]
        assert most_common_count == 1, f"{path.name}: duplicate section_no present"
        max_heading_len = max(len(s.heading) for s in sections)
        assert max_heading_len <= 300, f"{path.name}: heading exceeds column width"
        max_division_len = max(len(s.division or "") for s in sections)
        assert max_division_len <= 300, f"{path.name}: division exceeds column width"
        parsed_form_terms = len([s for s in sections if "-F" in s.section_no])
        expected_form_terms = _expected_form_term_count(data)
        assert parsed_form_terms == expected_form_terms, (
            f"{path.name}: parsed {parsed_form_terms} form terms, "
            f"independent scan expected {expected_form_terms}"
        )

    first_version_nos = {s.section_no for s in parse_docx(cached[0].read_bytes())}
    assert {"S1-F4-T1", "S1-F5-T1"} <= first_version_nos

    newest_sections = parse_docx(cached[-1].read_bytes())
    newest_f1 = [s for s in newest_sections if s.section_no.startswith("S1-F1-")]
    newest_f2 = [s for s in newest_sections if s.section_no.startswith("S1-F2-")]
    assert len(newest_f1) >= 30
    assert len(newest_f2) >= 38


def test_title_prose_split_at_exactly_150_chars():
    """Pins both sides of the heading threshold: 150 titles, 151 proses."""
    title_150 = "T" * 150
    prose_151 = "P" * 151
    data = build_docx(
        [
            ("Heading - PART", "Schedule 1—Forms"),
            ("New Form Heading", "Form 1"),
            ("New Form Heading", "A form title"),
            (None, f"1.\t{title_150}"),
            (None, "Body of the titled term."),
            (None, f"2.\t{prose_151}"),
            (None, "Continuation of the prose term."),
        ]
    )
    by_no = {s.section_no: s for s in parse_docx(data)}
    assert by_no["S1-F1-T1"].heading == title_150
    assert by_no["S1-F1-T1"].body_text == "Body of the titled term."
    assert by_no["S1-F1-T2"].heading == ""
    assert by_no["S1-F1-T2"].body_text == f"{prose_151} Continuation of the prose term."
